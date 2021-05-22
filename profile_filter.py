import os
import docx
import re
from prettytable import PrettyTable
import time
import xlsxwriter
import math
#################

from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager , PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser


###################

class Profile_filter:
    def __init__(self,pattern,path):
        self.Profile_path = path
        self.pattern = pattern
        self.output_file =''
        self.Profile_Files_pdf = []
        self.Profile_Files_docx = []
        self.table = None
        self.worksheet = None
        self.workbook = None
        self.cell_format_2 = None
        self.cell_format_3 =None
        self.row = 0
        self.xl = True
        self.progress_count = 0

    def output_setup(self):
        table_head = ['Filename','Match %', 'Match_count_in_each_file', 'Phone number', 'E-mail']
        self.table = PrettyTable(table_head,hrules=1)
        date = f'{time.strftime("%Y-%m-%d-%H%M%S")}'
        if self.xl:
            if os.path.isdir(self.Profile_path):
                self.output_file_xl = self.Profile_path + f'\\resultxl_{date}.xlsx'
            else:
                self.output_file_xl = (self.Profile_path.split(os.path.basename(self.Profile_path))[0]) + f'\\resultxl_{date}.xlsx'
            self.workbook = xlsxwriter.Workbook(self.output_file_xl)
            self.worksheet = self.workbook.add_worksheet()
            cell_format_1 = self.workbook.add_format({'bold':True, 'bg_color': '#C0C0C0', 'border':1})
            self.cell_format_2 =self.workbook.add_format({'bg_color': '#C4D79B', 'border':1})
            self.cell_format_3 =self.workbook.add_format({'bg_color': '#FFCC99', 'border':1})
            for column in range(len(table_head)):
                self.worksheet.write(self.row, column, table_head[column], cell_format_1)
            self.row += 1

    def filter_filenames(self):
        if os.path.isdir(self.Profile_path):
            for filename in os.listdir(self.Profile_path):
                if filename.endswith('.pdf') and not filename.startswith('~'):
                    self.Profile_Files_pdf.append(filename)
                if filename.endswith('.docx') and not filename.startswith('~'):
                    self.Profile_Files_docx.append(filename)
        if os.path.isfile(self.Profile_path):
            filename = self.Profile_path
            if filename.endswith('.pdf'):
                self.Profile_Files_pdf.append(filename)
            if filename.endswith('.docx'):
                self.Profile_Files_docx.append(filename)

    def read_docx(self,filename):
        pattern_count = [0] * len(self.pattern)
        hits = ''
        match = 0
        num = ''
        mail =''
        try:
            f = open(filename, 'rb')
            doc = docx.Document(f)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            for item in fullText:
                email = re.findall(r'[\w\.-]+@[\w\.-]+', item)
                phone = re.findall("\+?\d?\d?\-?\s?\d{3}[-.]?\d{3}[-.]?\d{4}", item)
                for num_match in phone:
                    num += num_match + " "
                for mail_match in email:
                    mail += mail_match + " "
                for index, pat in enumerate(self.pattern):
                    if re.search(re.escape(pat), item.lower()):
                        pattern_count[index] += 1
            f.close()
            for index,pat in enumerate(self.pattern):
                if pattern_count[index]>0:
                    hits = f'{hits} {pat}:{pattern_count[index]}'
                    match += 1
            match_percent = match/len(self.pattern) * 100
        except Exception as e:
            print(f"Error occured in {filename} ", e)
            match_percent = math.nan
            hits = 'Cannot read file'
        final_out = [filename.split('\\')[-1], f'{match_percent:.2f}', hits, num, mail]
        self.write_output(match_percent, final_out)
        self.progress_count += 1

    def read_pdf(self, filename):
        pattern_count = [0] * len(self.pattern)
        hits = ''
        match = 0
        num = ''
        mail = ''
        try:
            output_string = StringIO()
            with open(filename, 'rb') as in_file:
                parser = PDFParser(in_file)
                doc = PDFDocument(parser)
                rsrcmgr = PDFResourceManager()
                device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                for page in PDFPage.create_pages(doc):
                    interpreter.process_page(page)
                fullText = output_string.getvalue()
                email = re.findall(r'[\w\.-]+@[\w\.-]+', fullText)
                phone = re.findall("\+?\d?\d?\-?\s?\d{3}[-.]?\d{3}[-.]?\d{4}", fullText)
                for num_match in phone:
                    num += num_match + " "
                for mail_match in email:
                    mail += mail_match + " "
                for index, pat in enumerate(self.pattern):
                    if re.search(re.escape(pat), fullText.lower()):
                        pat_list = re.findall(re.escape(pat), fullText.lower())
                        pattern_count[index] += len(pat_list)
            for index,pat in enumerate(self.pattern):
                if pattern_count[index]>0:
                    hits = f'{hits} {pat}:{pattern_count[index]}'
                    match += 1
            match_percent = match/len(self.pattern) * 100
        except Exception as e:
            print(f"Error occured in {filename} ", e)
            match_percent = math.nan
            hits = 'Cannot read file'
        final_out = [filename.split('\\')[-1], f'{match_percent:.2f}', hits, num, mail]
        self.write_output(match_percent, final_out)
        self.progress_count += 1

    def write_output(self, match_percent, final_out):
        self.table.add_row(final_out)
        if self.xl:
            for column in range(5):
                if match_percent < 50:
                    c_fromat = self.cell_format_3
                else:
                    c_fromat = self.cell_format_2
                self.worksheet.write(self.row, column, final_out[column], c_fromat)
            self.row+=1

    def run_search(self):
        for profiles in self.Profile_Files_docx:
            print(f"Reading file: {profiles}")
            if os.path.isfile(profiles):
                self.read_docx(profiles)
            else:
                self.read_docx(self.Profile_path + '\\' + profiles)

        for profiles in self.Profile_Files_pdf:
            print(f"Reading file: {profiles}")
            if os.path.isfile(profiles):
                self.read_pdf(profiles)
            else:
                self.read_pdf(self.Profile_path + '\\' + profiles)

        if self.xl:
            self.workbook.close()

        print("Search completed.......")

    def start_profile_filter(self,filter_file = True):
        self.output_setup()
        if filter_file:
            self.filter_filenames()
        self.run_search()
        self.table.sortby = 'Match %'
        self.table.reversesort = True




